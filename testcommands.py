import datetime
import sqlite3
import os
import docx
import test
import discord
import datetime
from discord.ext import commands  # Again, we need this imported
import random
import asyncio


class TestCommands(commands.Cog):
    """Commands for test"""
    client = discord.Client()

    def __init__(self, bot: commands.Bot):
        self.bot = bot
        self.last_msg = None
        self.custom_prefix = bot.command_prefix
        self.bot.participating_in_test_right_now = []
        self.bot.number_of_tests = 0
        self.bot.number_of_test_answers = 0
        self.bot.number_of_test_answers_sets = 0
        self.bot.new_test = []
        self.bot.test_answers = []  # student #0 answers , student #1 answers etc.every element represents a test that a student wrote.
        self.bot.test_answer_sets = []  # list of test_answers objects.every element represents a test and the answers of all the students in that test.

    @commands.command(name="loadtest")
    async def loadtest(self, ctx, docname=None):
        """Φορτώνει το τεστ στη μνήμη"""
        try:
            if docname != None:
                test_loaded = False
                for i in range(len(self.bot.new_test)):
                    if self.bot.new_test[i].name_of_test == docname:
                        await ctx.send("Το τεστ είναι ήδη φορτωμένο στη μνήμη. ")
                        test_loaded = True
                        break
                if (test_loaded):
                    return

                if ctx.author.guild_permissions.administrator:
                    # We specify the document path here .
                    # If the document is in the file only the name is needed.
                    try:
                        doc = docx.Document(docname)
                    except Exception as error:
                        await ctx.author.send(f"Oops! {error}")
                        return
                    self.bot.number_of_tests += 1
                    index = self.bot.number_of_tests - 1
                    this_test = test.Test(docname)
                    self.bot.new_test.append(this_test)
                    question_objects = []
                    argument_probable_answers = []
                    for i in range(len(doc.paragraphs)):
                        question = doc.paragraphs[i].text.split("$")
                        if question[0] == "0":
                            question = doc.paragraphs[i].text.split("$", maxsplit=1)
                            temp_question_object = test.Question_type_0(docname, f"{i}", "0", f"{question[1]}")
                            self.bot.new_test[index].add_question_0(temp_question_object)


                        elif question[0] == "1":
                            for j in range(len(question) - 3):
                                argument_probable_answers.append(question[j + 3])
                            temp_question_object = test.Question_type_1(docname, f"{i}", "1", question[2],
                                                                        argument_probable_answers, question[1])
                            self.bot.new_test[index].add_question_1(temp_question_object)
                            argument_probable_answers = []


                        elif question[0] == "2":
                            for j in range(len(question) - 3):
                                argument_probable_answers.append(question[j + 3])
                            temp_question_object = test.Question_type_2(docname, f"{i}", "2", question[2],
                                                                        argument_probable_answers, question[1])
                            self.bot.new_test[index].add_question_2(temp_question_object)
                            argument_probable_answers = []
                        question_objects.append(temp_question_object)

                    self.bot.new_test[index].add_question_objects(question_objects)



                else:
                    message = await ctx.channel.send("Ουπς ! Μάλλον δεν έχετε δικαιώματα διαχειριστή.")
                    await asyncio.sleep(2)
                    await message.delete()
                    return
                await ctx.author.send("Η φόρτωση του τεστ στη μνήμη ήταν επιτυχής.")
            else:
                await ctx.author.send(f" ``Τα ορίσματα δεν προσφέρθηκαν σωστά``\n"
                                      f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}loadtest <name_of_test.docx>")


        except:
            await ctx.author.send(
                "Η φόρτωση του τεστ απέτυχε,η εντολή πρέπει να εκτελεστεί σε κανάλι σέρβερ από ένα διαχειριστή.")

    @commands.command(name="teststart")
    async def teststart(self, ctx, docname=None):
        "Ξεκινάει το τεστ που έχει φορτωθεί με την εντολή loadtest"

        if docname != None:
            test_not_loaded = True
            for i in range(len(self.bot.new_test)):
                if self.bot.new_test[i].name_of_test == docname:
                    # await ctx.send("test has already been loaded")
                    test_not_loaded = False
                    this_test = self.bot.new_test[i]
                    break
            if (test_not_loaded):
                await ctx.author.send("Αυτό το τεστ δεν έχει φορτωθεί στη μνήμη")
                return
            if ctx.author.id not in self.bot.participating_in_test_right_now:
                self.bot.participating_in_test_right_now.append(ctx.author.id)
            else:
                await ctx.author.send("Συμμετέχεις ήδη σε ένα τέστ . Ολοκλήρωσέ το πριν αρχίσεις άλλο")
                return
            if self.bot.banned_words == True: #during the test we don't want a filter .
                self.bot.banned_words = False
            self.bot.number_of_test_answers += 1
            index_student = self.bot.number_of_test_answers - 1
            this_student = test.Student_answers(ctx.author.id)
            self.bot.test_answers.append(this_student)

            def check(msg):
                return msg.author == ctx.author

            answer_objects = []
            for j in range(len(this_test.question_objects)):
                question = this_test.question_objects[j]
                question_type = question.get_type()
                question_main_body = question.get_main_body()

                def check_if_valid(answer, length):
                    try:
                        int(answer)  # answer is an integer
                        if ((int(answer) > length) or (int(answer) < 0)):
                            return 3  # answer exceeds boundaries
                        else:
                            return 0  # answer is ok
                    except:
                        try:
                            float(answer)  # answer is a float
                            return 1
                        except:
                            return 2  # answer is a string

                def check_if_splitted(answer):
                    try:
                        splitted = answer.split(",")
                        if len(splitted) < 1:
                            return False
                        elif splitted[0]=="" or splitted[-1]=="":
                            return False
                        else:
                            return True
                    except:
                        return False

                def inform_user_for_wrong_input(type_of_input):
                    if type_of_input == 3:
                        return "Η τιμή ξεπερνάει τα όρια . Προσπαθήστε ξανά :"
                    elif type_of_input == 2:
                        return "Προσφέρατε δεδομένο τύπου string , όχι integer. Προσπαθήστε ξανά :"
                    elif type_of_input == 1:
                        return "Προσφέρατε δεδομένο τύπου float , όχι integer . Προσπαθήστε ξανά :"
                    else:
                        return "Μη αναγνωρίσιμος τύπος input"

                def check_absence_of_duplicate(answer_tokens):
                    try:
                        for i in range(0, len(answer_tokens)):
                            for j in range(i + 1, len(answer_tokens)):
                                if int(answer_tokens[i]) == int(answer_tokens[j]):
                                    return False
                        return True
                    except Exception as error:
                        #print(error)
                        return False

                if question_type == "0":
                    embed = discord.Embed(title=f"```{question_main_body}```", description=f"Ερώτηση {j} (Ανάπτυξης)",
                                          colour=0x00dffa, timestamp=datetime.datetime.now())
                    embed.set_footer(
                        text=f"Παρακαλώ γράψτε την απάντησή σας\n(Πιέστε shift+enter για να αλλάξετε γραμμή)")
                    await ctx.author.send(embed=embed)




                    answer_permission = False
                    while (answer_permission==False):

                        answer = await self.bot.wait_for("message", check=check)
                        answer_string = str(answer.content)

                        while answer_string.startswith(f"{self.custom_prefix}"):
                            await ctx.author.send(
                                "Παρακαλώ μην γράφετε εντολές κατα τη διάρκεια ερωτήσεως ανάπτυξης.Προσπαθήστε ξανά")
                            answer = await self.bot.wait_for("message", check=check)
                            answer_string = str(answer.content)
                        perm_answer_str=""

                        while (perm_answer_str!= "nai" and perm_answer_str!="yes" and
                               perm_answer_str!="no" and perm_answer_str!="ναι" and
                               perm_answer_str!="όχι" and perm_answer_str!="οχι" and perm_answer_str!="oxi"):
                            await ctx.author.send("Είναι αυτή η τελική σας απάντηση ? \nΔεκτές απαντήσεις : **(ναι/όχι)**")
                            permission_answer = await self.bot.wait_for("message", check=check)
                            perm_answer_str = str(permission_answer.content)

                        if (perm_answer_str =="ναι") or (perm_answer_str=="yes") or (perm_answer_str=="nai"):
                            answer_permission=True

                        else:
                            await ctx.author.send("Παρακαλώ ξαναγράψτε την απάντησή σας")









                    temp_answer_object = test.Answer_type_0(docname, ctx.author.id, f"{j}", "0", answer_string)
                    self.bot.test_answers[index_student].add_answer_0(temp_answer_object)
                elif question_type == "1":

                    probable_answers = question.get_probable_answers()
                    probable_answers_test = ""
                    for x in range(len(probable_answers)):
                        probable_answers_test = probable_answers_test + f"\n{x}-{probable_answers[x]}"
                    embed = discord.Embed(title=f"```{question_main_body}```",
                                          description=f"Ερώτηση {j} (Μονής Επιλογής)",
                                          colour=0xfaef00, timestamp=datetime.datetime.now())
                    embed.set_footer(
                        text=f"Γράψε μόνο τον αριθμό της επιλογής σας,μεταξύ 0 και {(len(probable_answers) - 1)}")
                    embed.add_field(name="``Παρακαλώ επιλέξτε ένα από τα παρακάτω``", value=f"{probable_answers_test}")
                    await ctx.author.send(embed=embed)

                    answer = await self.bot.wait_for("message", check=check)
                    type_of_input = check_if_valid(str(answer.content), (len(probable_answers) - 1))
                    while type_of_input != 0:
                        await ctx.author.send(f"{inform_user_for_wrong_input(type_of_input)}")
                        answer = await self.bot.wait_for("message", check=check)
                        type_of_input = check_if_valid(str(answer.content), (len(probable_answers) - 1))

                    answer_int = int(answer.content)
                    temp_answer_object = test.Answer_type_1(docname, ctx.author.id, f"{j}", "1", answer_int)
                    self.bot.test_answers[index_student].add_answer_1(temp_answer_object)
                elif question_type == "2":
                    number_of_correct_answers = question.get_number_of_correct_answers()
                    probable_answers = question.get_probable_answers()
                    probable_answers_test = ""
                    for x in range(len(probable_answers)):
                        probable_answers_test = probable_answers_test + f"\n{x}-{probable_answers[x]}"
                    embed = discord.Embed(title=f"```{question_main_body}```",
                                          description=f"Ερώτηση {j} (Πολλαπλής Επιλογής)",
                                          colour=0xfa7400, timestamp=datetime.datetime.now())
                    embed.set_footer(
                        text=f"Παρακαλώ γράψτε τους αριθμούς που επιλέξατε μεταξύ 0 και {(len(probable_answers) - 1)}"
                             f"χωρισμένους με κόμμα . Για παράδειγμα : 2,6,5")
                    embed.add_field(name=f"``Επιλέξτε {number_of_correct_answers} από τις παρακάτω επιλογές :``",
                                    value=f"{probable_answers_test}")
                    await ctx.author.send(embed=embed)

                    answer = await self.bot.wait_for("message", check=check)
                    answer_string = str(answer.content)
                    is_splitted = check_if_splitted(answer_string)
                    number_of_tokens_ok = True
                    type_of_tokens_ok = True

                    if is_splitted:
                        answer_tokens = answer_string.split(",")
                        if len(answer_tokens) == number_of_correct_answers:
                            for answer_token in answer_tokens:
                                type_of_input = check_if_valid(answer_token, (len(probable_answers) - 1))
                                if type_of_input != 0:
                                    type_of_tokens_ok = False
                                    break
                        else:
                            number_of_tokens_ok = False
                        absence_of_duplicate = check_absence_of_duplicate(answer_tokens)

                    while (is_splitted == False) or (type_of_tokens_ok == False) or (number_of_tokens_ok == False) or (
                            absence_of_duplicate == False):
                        try:
                            if is_splitted == False:
                                await ctx.author.send(
                                    "Ή δεν χωρίσατε τίποτα με κόμμα ή δεν προσφέρατε των απαιτούμενο αριθμό απαντήσεων\n"
                                    "Σιγουρευτείτε ότι δεν αφήνετε κενά ανάμεσα στα κόμματα .Προσπαθήστε ξανά.")
                            elif number_of_tokens_ok == False:
                                await ctx.author.send(
                                    f"**Δεν προσφέρατε τον αναμενόμενο αριθμό απαντήσεων. Προσπαθήστε ξανά.**\n"
                                    f"``Απαιτούμενος αριθμός απαντήσεων : {number_of_correct_answers}``\n"
                                    f"``Προσφερόμενες απαντήσεις χρήστη : {len(answer_tokens)}``\n")
                            elif absence_of_duplicate == False:
                                await ctx.author.send("Έχετε εισάγει ένα όρισμα πολλαπλές φορές. Προσπαθήστε ξανά.")
                            else:
                                await ctx.author.send(f"{inform_user_for_wrong_input(type_of_input)}")
                            type_of_tokens_ok = True
                            number_of_tokens_ok = True
                            answer = await self.bot.wait_for("message", check=check)
                            answer_string = str(answer.content)
                            is_splitted = check_if_splitted(answer_string)
                            if is_splitted:

                                answer_tokens = answer_string.split(",")
                                if len(answer_tokens) == number_of_correct_answers:

                                    for answer_token in answer_tokens:
                                        type_of_input = check_if_valid(answer_token, (len(probable_answers) - 1))

                                        if type_of_input != 0:
                                            type_of_tokens_ok = False
                                            break
                                else:
                                    number_of_tokens_ok = False
                                absence_of_duplicate = check_absence_of_duplicate(answer_tokens)
                        except Exception as error:
                            await ctx.author.send(f"Κάτι πήγε στραβά . Error : {error}. \n"
                                                  f"Κανένας λόγος άγχους . Δεν πειράζει . Προσπάθησε ξανά. ")
                            continue

                    temp_answer_object = test.Answer_type_2(docname, ctx.author.id, f"{j}", "2", answer_string)
                    self.bot.test_answers[index_student].add_answer_2(temp_answer_object)
                answer_objects.append(temp_answer_object)
            await ctx.author.send("Πολύ καλά! Μόλις ολοκληρώσατε την εξέτασή σας . Καλή συνέχεια!")
            for i in range(len(self.bot.participating_in_test_right_now)):
                if self.bot.participating_in_test_right_now[i] == ctx.author.id:
                    self.bot.participating_in_test_right_now.pop(i)
                    break
            self.bot.test_answers[index_student].add_answer_objects(answer_objects)
            answer_set_not_found = True
            for o in range(len(self.bot.test_answer_sets)):
                name_of_answer_set = self.bot.test_answer_sets[o].get_name_of_test()
                if name_of_answer_set == docname:
                    answer_set_not_found = False
                    break
            if (answer_set_not_found):
                answer_set = test.Test_answers(docname)
                self.bot.test_answer_sets.append(answer_set)
            else:
                answer_set = self.bot.test_answer_sets[o]
            answer_set.add_student_answers(self.bot.test_answers[index_student])


        else:
            await ctx.author.send(f" ``Τα ορίσματα δεν προσφέρθηκαν σωστά``\n"
                                  f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}start <name_of_test.docx>")

    @commands.command(name="participants")
    async def participants(self, ctx):
        """Δείχνει όσους συμμετέχουν στο τεστ αυτή τη στιγμή"""
        try:
            if ctx.author.guild_permissions.administrator:
                users = ""
                for element in self.bot.participating_in_test_right_now:
                    user = await self.bot.fetch_user(element)
                    users+=f"[{user.name}]"
                if users=="":
                    await ctx.author.send("Κανένας δεν συμμετέχει αυτή τη στιγμή")
                else:
                    await ctx.author.send(f"{users}")
            else:
                await ctx.author.send("Η εντολή αυτή χρησιμοποιείται μόνο από διαχειριστές σε κανάλι σέρβερ")
        except:
            if ctx.guild == None:
                await ctx.author.send("Δεν βρέθηκε κάποιος σέρβερ.Η εντολή πρέπει να πληκτρολογηθεί\n"
                                      "σε ειδικό κανάλι σέρβερ διαμορφωμένο για την λειτουργία agentchat.")
            return

    @commands.command(name="showme")
    async def showme(self, ctx, name_of_test=None, name_of_student=None):
        """ Δείχνει τις απαντήσεις ενός συγκεκριμένου χρήστη.
            <command_prefix>showme <name_of_test.docx> <student_id>
        """
        try:
            if (name_of_test != None) and (name_of_student != None) and (ctx.author.guild_permissions.administrator):
                answer_set_not_found = True
                for o in range(len(self.bot.test_answer_sets)):
                    name_of_answer_set = self.bot.test_answer_sets[o].get_name_of_test()
                    if name_of_answer_set == name_of_test:
                        answer_set_not_found = False
                        break
                if (answer_set_not_found):
                    await ctx.author.send("Το αρχείο δε βρέθηκε")
                    return
                else:
                    answer_set = self.bot.test_answer_sets[o]
                    student_answers = answer_set.get_student_answers()
                    student_name_not_found = True
                    for i in range(len(student_answers) - 1, -1, -1):
                        student_name = student_answers[i].get_student_name()
                        if str(student_name) == name_of_student:
                            student_name_not_found = False
                            break
                    if (student_name_not_found):
                        await ctx.author.send("Το αρχείο δε βρέθηκε")
                        return
                    else:
                        all_the_answers = student_answers[i].get_answer_objects()
                        answer_str=""
                        for i in range(len(all_the_answers)):
                            # this_type = all_the_answers[i].get_type()
                            answer = all_the_answers[i].get_answer_body()
                            answer_str+=f"\n{answer}"
                        await ctx.author.send(f"{answer_str}")

            else:
                await ctx.author.send(f"``Τα ορίσματα δεν προσφέρθηκαν σωστά.``\n"
                                    f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}start <name_of_test.docx>\n"
                                      f"Μόνο διαχειριστές έχουν πρόσβαση σε αυτή την εντολή")
        except Exception as error:
            await ctx.author.send(f"{error}")


    @commands.command(name="stats")
    async def stats(self, ctx, name_of_test=None):
        """<command_prefix>stats <name_of_test.docx>"""

        try:
            if ctx.author.guild_permissions.administrator:
                if (name_of_test != None):
                    answer_set_not_found = True
                    for o in range(len(self.bot.test_answer_sets)):
                        name_of_answer_set = self.bot.test_answer_sets[o].get_name_of_test()
                        if name_of_answer_set == name_of_test:
                            answer_set_not_found = False
                            break
                    if (answer_set_not_found):
                        await ctx.author.send("Το αρχείο δε βρέθηκε")
                        return
                    else:
                        answer_set = self.bot.test_answer_sets[o]
                        student_answers = answer_set.get_student_answers()

                        # this is a procedure to remove the duplicate answers by students ,
                        # and take into consideration only their last attempt of the test.
                        list_of_names = []
                        for i in range(len(student_answers) - 1, -1, -1):
                            student_name = student_answers[i].get_student_name()
                            if student_name in list_of_names:
                                student_answers.pop(i)

                            else:
                                list_of_names.append(student_name)

                        number_of_answers = len(student_answers[0].get_types_by_position())

                        list_of_questions = []
                        list_of_answers = []
                        for j in range(number_of_answers):  # For <j> question .
                            type_of_answer = student_answers[0].get_type_by_position(j)
                            if type_of_answer == "0":
                                continue
                            if len(list_of_answers) > 0:

                                list_of_answers.clear()

                            for i in range(len(student_answers)):  # what did <i> student answer .

                                if type_of_answer == "1":
                                    answer_object_1 = student_answers[i].get_answer_object(j)
                                    element = answer_object_1.get_answer_body()

                                    list_of_answers.append(element)


                                if type_of_answer == "2":
                                    answer_object_2 = student_answers[i].get_answer_object(j)
                                    elements = answer_object_2.get_answer_body()
                                    for element_in_list in elements:

                                        list_of_answers.append(element_in_list)

                            list_immutable = tuple(list_of_answers)
                            list_of_questions.append(list_immutable)

                        #print(list_of_questions)
                        from collections import Counter
                        embed = discord.Embed(title=f"```Στατιστικά για {name_of_test}```",
                                            colour=0xe207ff, timestamp=datetime.datetime.now())
                        for i in range(len(list_of_questions)):
                            counter = Counter(list_of_questions[i])
                            counter_most_common = counter.most_common()

                            stats_message = (f"")
                            for j in counter_most_common:
                                stats_message += (
                                    f"Η επιλογή {j[0]} επιλέχθηκε από {j[1]} άτομα ({j[1] / len(list_of_names) * 100} %)\n")
                            embed.add_field(name=f"Ερώτηση {i}", value=f"{stats_message}", inline=False)
                        await ctx.channel.send(embed=embed)

                else:
                    await ctx.send(f"``Τα ορίσματα δεν προσφέρθηκαν σωστά.``\n"
                                f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}start <name_of_test.docx>")
            else:
                await ctx.author.send("Η εντολή αυτή χρησιμοποιείται μόνο από διαχειριστές σε κανάλι σέρβερ")
        except:
            await ctx.author.send("Η εμφάνιση των στατιστικών απέτυχε.")

    @commands.command(name="printanswers")
    async def printanswers(self, ctx, name_of_test=None):
        """
            Δημιουργεί όλα τα αρχεία χρηστών με τις απαντήσεις τους.
            <command_prefix>printanswers <name_of_test.docx>
        """
        try:
            if ctx.author.guild_permissions.administrator:
                if (name_of_test != None):
                    answer_set_not_found = True
                    questions_set_not_found = True
                    for o in range(len(self.bot.test_answer_sets)):
                        name_of_answer_set = self.bot.test_answer_sets[o].get_name_of_test()
                        if name_of_answer_set == name_of_test:
                            answer_set_not_found = False
                            break
                    if (answer_set_not_found):
                        await ctx.author.send("Το σετ των απαντήσεων δε βρέθηκε")
                        return

                    for p in range(len(self.bot.new_test)):
                        name_of_question_set = self.bot.new_test[p].get_name_of_test()
                        if name_of_question_set == name_of_test:
                            questions_set_not_found = False
                            break

                    if (questions_set_not_found):
                        await ctx.author.send("Το σετ των ερωτήσεων δεν βρέθηκε")
                        return
                    else:
                        answer_set = self.bot.test_answer_sets[o]
                        question_set = self.bot.new_test[p]
                        student_answers = answer_set.get_student_answers()
                        # this is a procedure to remove the duplicate answers by students ,
                        # and take into consideration only their last attempt of the test.
                        list_of_names = []
                        for i in range(len(student_answers) - 1, -1, -1):
                            student_name = student_answers[i].get_student_name()
                            if student_name in list_of_names:
                                student_answers.pop(i)

                            else:
                                list_of_names.append(student_name)

                        mycwd = os.getcwd()

                        name_of_test_directory = name_of_test.split('.')[0]
                        if not os.path.exists(f"{name_of_test_directory}"):
                            os.mkdir(f"{name_of_test_directory}.")
                            await ctx.author.send(f"Folder {name_of_test_directory} with answers created in : {mycwd}")

                        try:
                            users_not_in_database = ""
                            for i in range(len(student_answers) - 1, -1, -1):
                                student_name = student_answers[i].get_student_name()
                                this_student_doc = docx.Document()
                                student_in_database = False
                                try:
                                    connection = sqlite3.connect("studentsdata.db")
                                    cursor = connection.cursor()
                                    finduser = f"SELECT * FROM students WHERE DISCORD_ID = '{str(student_name)}'"
                                    cursor.execute(finduser)
                                    result = cursor.fetchall()
                                    connection.close()
                                    for l in range(len(result[0])):
                                        if l == 0:
                                            user_discord_id = f"{str(result[0][l])}"
                                        elif l == 1:
                                            user_name = f"{str(result[0][l])}"
                                        elif l == 2:
                                            user_surname = f"{str(result[0][l])}"
                                        elif l == 3:
                                            user_aem = f"{str(result[0][l])}"
                                        elif l == 4:
                                            user_type = f"{str(result[0][l])}"
                                        elif l == 5:
                                            user_mail = f"{str(result[0][l])}"
                                        else:
                                            break
                                    student_in_database = True

                                except:
                                    users_not_in_database+=f"[ ** {str(student_name)} ** ]"

                                if (student_in_database):

                                    this_student_doc.add_heading(f"{user_aem} {user_name} {user_surname}")
                                    this_student_doc.add_paragraph(f"discord_id = {user_discord_id}\n"
                                                                   f"email = {user_mail}")
                                else:
                                    this_student_doc.add_heading(f"{student_name}")
                                test_questions = question_set.get_question_objects()
                                this_student_answers = student_answers[i].get_answer_objects()
                                for u in range(len(this_student_answers)):

                                    this_question_main_body = test_questions[u].get_main_body()
                                    this_question_type = test_questions[u].get_type()
                                    this_answer = this_student_answers[u].get_answer_body()
                                    if this_question_type !="0":
                                        this_question_probable_answers = test_questions[u].get_probable_answers()
                                        this_student_doc.add_paragraph(f"Ερώτηση {u} :\n{this_question_main_body}"
                                                                       f"\nΠιθανές απαντήσεις : {this_question_probable_answers}")
                                        this_question_correct_answers = test_questions[u].get_correct_answers()
                                        sum = 0
                                        if this_question_type == "2":
                                            for t in range(len(this_question_correct_answers)):
                                                if this_question_correct_answers[t] in this_answer:
                                                    sum+=1
                                            this_question_score = f"{sum}/{len(this_question_correct_answers)}"
                                            answer_string = []
                                            for t in range(len(this_answer)):
                                                answer_string.append(str(this_question_probable_answers[this_answer[t]]))
                                        elif this_question_type == "1":
                                            if this_question_correct_answers == this_answer:
                                                sum+=1
                                            this_question_score = f"{sum}/1"
                                            answer_string = str(this_question_probable_answers[this_answer])
                                        this_student_doc.add_paragraph(f"Απάντηση μαθητή  : {answer_string}\n"
                                                                       f"Βαθμός ερώτησης : {this_question_score}")
                                    elif this_question_type =="0":
                                        this_student_doc.add_paragraph(f"Ερώτηση {u} :\n{this_question_main_body}")
                                        this_student_doc.add_paragraph(f"Απάντηση μαθητή  : {this_answer}")
                                    this_student_doc.add_paragraph(f"=============================================")
                                os.chdir(
                                    f"{name_of_test_directory}")  # changes current directory to that of the folder we are going to store the answers
                                if(student_in_database):
                                    this_student_doc.save(f"{user_aem} {user_name} {user_surname}.docx")
                                else:
                                    this_student_doc.save(f"{student_name}.docx")
                                os.chdir(f"{mycwd}") #changes current directory to the default local directory.
                        except Exception as error:
                            await ctx.author.send("η σύνδεση με τη βάση δεδομένων χρηστών απέτυχε.\n"
                                                  f"ERROR : {error}")
                        os.chdir(f"{mycwd}")
                        if users_not_in_database!="":
                            await ctx.author.send(f"```Οι χρήστες με τα παρακάτω discord id ίσως δεν έχουν κάνει εγγραφή:```\n{users_not_in_database}\n"
                                                  f"\n``Πατήστε <@id> σε κάποιο κανάλι για να σας δείξει \nποιοι είναι,όπου id κάποιο id από τα παραπάνω.``")
                else:
                    await ctx.send(f"``Τα ορίσματα δεν προσφέρθηκαν σωστά.``\n"
                                   f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}start <name_of_test.docx>")
            else:
                await ctx.author.send("Η εντολή αυτή χρησιμοποιείται μόνο από διαχειριστές σε κανάλι σέρβερ")
        except Exception as error:
            await ctx.author.send("Η δημιουργία των αρχείων των απαντήσεων απέτυχε.\n"
                                  f"ERROR : {error}")

    @commands.Cog.listener()
    async def on_ready(self):
        print('TestCommands cog ready')



def setup(bot: commands.Bot):
    bot.add_cog(TestCommands(bot))