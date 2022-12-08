
import sqlite3
import os
import docx
import discord
from discord.ext import commands  # Again, we need this imported
import random
import asyncio

class AgentChatCommands(commands.Cog):
    """Commands for agent"""
    client = discord.Client()

    def __init__(self, bot: commands.Bot):
        self.bot = bot
        self.last_msg = None
        self.custom_prefix = bot.command_prefix
        self.bot.agent_questions = {}
        self.bot.agent_counters = {}
        self.bot.agent_teams = {}
        self.bot.agent_teams["teams"] = {}
        self.bot.agent_teams["ids"] = {}
        self.bot.agent_students = {}
        self.bot.agent_answers = {}
        # this list prevents students from activating the agentstart command when already in an agent chat activity.
        self.bot.participating_in_agent_right_now = []

    """
    ----->These are some examples of how dictionaries are structured for agent chat.<------
    
    First dictionary is the data of the students (if there are any)
    
    agent_students = { discord_id1 :  { 'name' : name1 , 'surname' : name2 , 'mail' : mail1  , data ... etc } ,
                        discord_id2 : {'name' : name2  , 'surname' : name3 ,  'aem' : aem2  .... etch }}   
    """ #agent_students
    """
    This is a hashtable of counters . Values represent how many qustions a keyword holds.
    Later we can use this to create key-value pairs into the agents dictionary.

    agent_counters = {"agent1.docx" : {"keyword1" : value1,
                            "keyword2" : value2},
                "agent2.docx" : {"keyword1" : value1,
                            "keyword2" : value2,
                            "keyword3" : value3}}
    """ #agent_counters
    """    
    This is an example of how the dictionary is being structured where we store the questions.
    Every keyword represents a trigger. Then the student is called to answer one of the questions.
    For example, keyword1 holds four questions(0-3) . First element of the list represents the type of the
    question and the second element is the main body of the question.
                        
    agent_questions = { "agent1.docx" : { "keyword1" : { 0 : [0,question_main_body_1],
                                                    1 : [0,question_main_body_2],
                                                    2 : [0,question_main_body_3],
                                                    3 : [1,question_main_body_4]},
                                     "keyword2" : { 0 : [0,question_main_body_1],
                                                    1 : [0,question_main_body_2]},
                                     "keyword3" : { 0 : [0,question_main_body_1],
                                                    1 : [0,question_main_body_2],
                                                    2 : [0,question_main_body_3]}},
                        "agent2.docx" :   {"keyword1" : { 0 : [0,question_main_body_1],
                                                    1 : [0,question_main_body_2],
                                                    2 : [0,question_main_body_3]}}}
    """ #agent_questions
    """                                           
    Of course we need to divide the students .
    Each student is assigned a unique team and a unique agent.
    Teams dictionary is divided into two nested dictionaries. One holds 
    the ids as values and the other the teams with the agents.
    Each team though holds a specific channel for the agent to work,so teams will
    be referenced with channel ids.agent_teams dictionary is translated as 
    participating_right_now dictionary . 
    
    !!! By any means you must NEVER use the same channel to start chat with two different agents!!! 
    
    agent_teams = { "teams" : { channel_id1 : [user_id1,user_id2],
                                channel_id2 : [user_id3,user_id4,user_id5]}
                    "ids"   : { user_id1 : [channel_id1,'agent1.docx']
                                user_id2 : [channel_id1,'agent1.docx']
                                user_id3 : [channel_id2,'agent1.docx']
                                user_id4 : [channel_id2,'agent1.docx']
                                user_id5 : [channel_id2,'agent1.docx'}}
    """ #agent_teams
    """
    And last but not least we need to store the answers of the students somewhere
    
    agent_answers = { student_id1 : { "agent1.docx" : {"keyword1" : { number_of_the_question : [type_of_question,answer_main_body],
                                                                  number_of_the_2nd_question : [type_of_question,answer_main_body]},
                                                    "keyword2" : {number_of_the_question : [type_of_question,answer_main_body]}},
                                        "agent2.docx"  : { "keyword1" : {nunmber_of_the_question : [type_of_question,answer_main_body]}}},
                        student_id2 : {..........}}
                
    """ #agent_answers

    @commands.command(name = "loadagent")
    async def loadagent(self,ctx,name_of_agent=None):
        """Load a scripted file with keywords in memory"""
        try:
            if name_of_agent != None:
                if name_of_agent in self.bot.agent_questions:
                    await ctx.author.send("Ο πράκτορας είναι ήδη φορτωμένος στη μνήμη")
                    return
                if ctx.author.guild_permissions.administrator:
                    # We specify the document path here .
                    # If the document is in the file only the name is needed.
                    try:
                        agent_doc = docx.Document(name_of_agent)
                    except Exception as error:
                        await ctx.author.send(f"Ουπς ! {error} " )
                        return
                    """ 
                        From now on we set nicknames so that code is shorter
                        agent counters dictionary is refered as acounters
                        agent_questions dictionary as aquestions
                        name_of_agent string is refered as aname
                    """
                    acounters = self.bot.agent_counters
                    aquestions= self.bot.agent_questions
                    aname = name_of_agent

                    """
                        Now we start the loading procedure
                        The document that is loaded should have this format:
                        
                        type_of_question$keyword$question_main_body
                        
                        For example : 
                        
                        0$sunny$Is the weather sunny today ?
                        
                        For this project only question type 0 is supported.
                        But the tools for implementing more types are provided.
                        Question type 0 is a question where the answer is a paragraph.
                        The student must express his/her opinion with words.
                                           
                    """

                    acounters[aname] = {}
                    aquestions[aname] = {}
                    for i in range(len(agent_doc.paragraphs)):
                        question = agent_doc.paragraphs[i].text.split("$")
                        keyword = question[1]
                        #print((question[2]))
                        if question[0] == "0":
                            if keyword not in aquestions[aname]:
                                aquestions[aname][keyword] ={}
                            aquestions[aname][keyword][len(aquestions[aname][keyword])] = [int(question[0]),question[2]]
                            acounters[aname][keyword]=len(aquestions[aname][keyword])

                    """
                    prints the questions dictionary 
                    
                    for key,value in aquestions[aname].items():
                        print(key)
                        for value_key,value_value in value.items():
                            print("    ", value_key,value_value)
                            
                    """
                else:
                    message = await ctx.channel.send("Ουπς ! Μάλλον δεν έχετε δικαιώματα διαχειριστή.")
                    await asyncio.sleep(2)
                    await message.delete()
                    return
                await ctx.author.send("Η φόρτωση του πράκτορα στη μνήμη ήταν επιτυχής.")
            else:
                await ctx.author.send(f" ``Τα ορίσματα δεν προσφέρθηκαν σωστά``\n"
                              f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}loadtest <name_of_test.docx>")
        except Exception as error:
            await ctx.author.send(f"Η φόρτωση του τεστ απέτυχε .\n {error}")

    @commands.command(name="agentstart")
    async def agentstart(self,ctx,name_of_agent=None):
        """Participate in agent chat with a team"""

        if name_of_agent != None:
            if ctx.author.id in self.bot.participating_in_agent_right_now:
                return
            if self.bot.agent == False:
                await ctx.author.send(f"Η λειτουργία πρακτόρων είναι κλειστή\n"
                                      f"Κάποιος διαχειριστής θα πρέπει να την ανοίξει\n "
                                      f"με την εντολή ``{self.custom_prefix}enable agent``")
                return
            if ctx.guild == None:
                await ctx.author.send("Δεν βρέθηκε κάποιος σέρβερ.Η εντολή πρέπει να πληκτρολογηθεί\n"
                                      "σε ειδικό κανάλι σέρβερ διαμορφωμένο για την λειτουργία agentchat.")
                return

            """ 
                From now on we set nicknames so that code is shorter
                agent_teams dictionary is refered as ateams
                agent_answers dictionary is refered as aanswers
                name_of_agent string is refered as aname
        
            """
            aname = name_of_agent
            ateams = self.bot.agent_teams
            aanswers = self.bot.agent_answers
            astudents = self.bot.agent_students
            def check(msg):
                return msg.author == ctx.author

            async def permission(ctx):

                perm_answer_str = "preparing for while loop.perm_answer_str initialization"
                while (perm_answer_str!= "nai" and perm_answer_str!="yes" and
                    perm_answer_str!="no" and perm_answer_str!="ναι" and
                    perm_answer_str!="όχι" and perm_answer_str!="οχι" and perm_answer_str!="oxi"):

                        permission_answer = await self.bot.wait_for("message", check=check)
                        perm_answer_str = str(permission_answer.content)
                        if perm_answer_str.startswith(f"{self.custom_prefix}"):
                            await ctx.author.send(
                                "Παρακαλώ μη γράφετε εντολές κατά τη διαδικασία επιβεβαίωσης.Προσπαθήστε ξανά\n"
                                "Δεκτές απαντήσεις : **(ναι/όχι)**")

                if (perm_answer_str == "nai") or (perm_answer_str=="ναι") or (perm_answer_str=="yes"):
                    return True
                else:
                    return False

            if aname not in self.bot.agent_questions:
                await ctx.author.send("Αυτός ο πράκτορας δεν έχει φορτωθεί στη μνήμη")
                return

            if self.bot.banned_words == True:  # during the test we don't want any filter .
                self.bot.banned_words = False
            if ctx.author.id not in astudents:
                try:
                    connection = sqlite3.connect("studentsdata.db")
                    cursor = connection.cursor()
                    finduser = f"SELECT * FROM students WHERE DISCORD_ID = '{str(ctx.author.id)}'"
                    cursor.execute(finduser)
                    result = cursor.fetchall()
                    connection.close()
                    for i in range(len(result[0])):
                        if i == 0:
                            user_discord_id = f"{str(result[0][i])}"
                        elif i == 1:
                            user_name = f"{str(result[0][i])}"
                        elif i == 2:
                            user_surname = f"{str(result[0][i])}"
                        elif i == 3:
                            user_aem = f"{str(result[0][i])}"
                        elif i == 4:
                            user_type = f"{str(result[0][i])}"
                        elif i == 5:
                            user_mail = f"{str(result[0][i])}"
                        else:
                            break
                    astudents[ctx.author.id] = {}
                    astudents[ctx.author.id]["id"] = user_discord_id
                    astudents[ctx.author.id]["name"] = user_name
                    astudents[ctx.author.id]["surname"] = user_surname
                    astudents[ctx.author.id]["mail"] = user_mail
                    astudents[ctx.author.id]["type"] = user_type
                    astudents[ctx.author.id]["aem"] = user_aem
                except Exception as error:
                    await ctx.author.send(f"Τα στοιχεία σας δεν βρέθηκαν στη βάση δεδομένων . Καλό είναι να κάνετε εγγραφή.\n"
                                          f"{error}")


            if ctx.author.id in ateams["ids"]:

                active_channel = ateams["ids"][ctx.author.id][0]
                current_team = ateams["teams"][active_channel]
                current_team_list = ""
                for element in current_team:
                    try:
                        current_team_list += f"{str(astudents[element]['name'])} {str(astudents[element]['surname'])}\n"
                    except:
                        current_team_list +=f"{str(element)}\n"
                await ctx.author.send(f"Είστε ήδη σε μια ομάδα .\n"
                                      f"Τωρινή ομάδα : {current_team_list}"
                                      f"Θέλετε να αλλάξετε ομάδα ;\n"
                                    f"Απαντήστε με **'ναι'** για να ακυρωθεί η συμμετοχή σας \n"
                                    "Απαντήστε με **'όχι'** για να παραμείνετε στην ομάδα σας και να συνεχίσετε τη συνομιλία\n"
                                      "Αν ακυρώσετε τη συμμετοχή σας μπορείτε και πάλι αργότερα να ξανασυμμετέχετε στον πράκτορα της επιλογής σας\n"
                                      "Δεκτές απαντήσεις : **(ναι/όχι)**")
                self.bot.participating_in_agent_right_now.append(ctx.author.id)
                permission_to_change_team = await permission(ctx)
                if permission_to_change_team:
                    await ctx.author.send(f"Θέλετε επίσης οι προηγούμενες απαντήσεις σας (αν είχατε) στον πράκτορα {aname} να σβηστούν;\n"
                                          f"Δεκτές απαντήσεις : **(ναι/όχι)**")
                    permission_to_erase_answers = await permission(ctx)
                    if permission_to_erase_answers:
                        if ctx.author.id in aanswers:
                            if aname in aanswers[ctx.author.id]:
                                aanswers[ctx.author.id].pop(aname)
                                await ctx.author.send(f"Οι προηγούμενες απαντήσεις για το πράκτορα {aname} έχουν διαγραφεί\n"
                                                      f"αν θες να συμμετάσχεις ξανά σε συνομιλία με πράκτορα απλώς ξαναπάτα τη σχετική εντολή\n"
                                                      f"{self.custom_prefix}agentstart <name_of_agent.docx>")
                            else:
                                await ctx.author.send("Δεν έχει καταγραφεί καμία απάντηση στο συγκεκριμένο πράκτορα  (^_^) \n"
                                                      "Κανένας λόγος ανησυχίας!!!,απλώς **ξαναπατήστε την εντολή agentstart μαζί με το όνομα του πράκτορα**\n"
                                                      "με τον οποίο θέλετε να ξεκινήσετε chat με κάποιο συνομιλιτή")

                        else:
                            await ctx.author.send("Μάλλον δεν έχετε καν συμμετάσχει σε κάποιο πράκτορα  (^_^) \n"
                                                  "Κανένας λόγος ανησυχίας!!!,απλώς **ξαναπατήστε την εντολή agentstart μαζί με το όνομα του πράκτορα**\n"
                                                  "με τον οποίο θέλετε να ξεκινήσετε chat με κάποιο συνομιλιτή")
                    else:
                        await ctx.author.send("Η συμμετοχή σας ακυρώθηκε . Μπορείτε να συμμετέχετε σε καινούριο ή τον ίδιο πράκτορα.\n"
                                        "Οι απαντήσεις σας από προηγούμενες προσπάθειες παραμένουν καταγραμμένες(αν είχατε) αλλά \n"
                                        "στην περίπτωση καινούριας προσπάθειας πάνω στις ίδιες ερωτήσεις οι καινούριες απαντήσεις θα πάρουν τη θέση τους.")

                    active_channel = ateams["ids"].pop(ctx.author.id)[0]
                    ateams["teams"][active_channel].remove(ctx.author.id)
                    if len(ateams["teams"][active_channel]) == 0:
                        ateams["teams"].pop(active_channel)

                else:
                    await ctx.author.send("Πολύ καλά. Παραμένετε στην ίδια ομάδα.")
                self.bot.participating_in_agent_right_now.remove(ctx.author.id)
            else:
                if ctx.channel.id in ateams["teams"]:
                    ausersid = ateams["teams"][ctx.channel.id][0]
                    auserschannel = ateams["ids"][ausersid][0]
                    ausersagent = ateams["ids"][ausersid][1]
                    if ctx.channel.id == auserschannel and aname!=ausersagent :
                        await ctx.author.send(f"Αυτό το κανάλι χρησιμοποιείται από το πράκτορα {ausersagent}.\n"
                                              f"H συμμετοχή σας λαμβάνεται ως άκυρη")
                        return


                ateams["ids"][ctx.author.id] = [ctx.channel.id,aname]
                if ctx.channel.id in ateams["teams"]:
                    ateams["teams"][ctx.channel.id].append(ctx.author.id)
                else:
                    ateams["teams"][ctx.channel.id] = [ctx.author.id]
                current_team = ateams["teams"][ctx.channel.id]
                current_team_list = ""
                for element in current_team:
                    try:
                        temp_team_list = f"{str(astudents[element]['name'])} {str(astudents[element]['surname'])}\n"
                    except:
                        temp_team_list = f"{str(element)}\n"
                    current_team_list+=temp_team_list
                await ctx.author.send(f"Πολύ καλά , ξεκινήσατε chat με τον πράκτορα {name_of_agent} στο κανάλι {ctx.channel.name} \n"
                                      f"Η ομάδα σας είναι η : {current_team_list}")
        else:
            await ctx.author.send(f" ``Τα ορίσματα δεν προσφέρθηκαν σωστά``\n"
                                    f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}agentstart <name_of_agent.docx>")

    @commands.command(name="agentfinish")
    async def agentfinish(self,ctx):
        try:
            ateams = self.bot.agent_teams
            active_channel = ateams["ids"].pop(ctx.author.id)[0]
            ateams["teams"][active_channel].remove(ctx.author.id)
            if len(ateams["teams"][active_channel])==0:
                ateams["teams"].pop(active_channel)
            await ctx.author.send("Αποχωρήσατε από την ομάδα σας")
        except:
            await ctx.author.send("Η διαγραφή απέτυχε. Μάλλον δεν συμμετήχατε σε κάποια ομάδα")

    @commands.Cog.listener()
    async def on_message(self, ctx):
        try:
            if ctx.channel.id in self.bot.participating_in_agent_right_now:
                return
            def check(msg):
                return msg.author == ctx.author
            async def permission(ctx):

                perm_answer_str = "preparing for while loop.perm_answer_str initialization"
                while (perm_answer_str != "nai" and perm_answer_str != "yes" and
                    perm_answer_str != "no" and perm_answer_str != "ναι" and
                    perm_answer_str != "όχι" and perm_answer_str != "οχι" and perm_answer_str != "oxi"):
                    if (perm_answer_str != "nai" and perm_answer_str != "yes" and
                    perm_answer_str != "no" and perm_answer_str != "ναι" and
                    perm_answer_str != "όχι" and perm_answer_str != "οχι" and perm_answer_str != "oxi"):
                        await ctx.channel.send("Δεκτές απαντήσεις : **(ναι/όχι)**")

                    permission_answer = await self.bot.wait_for("message", check=check)
                    perm_answer_str = str(permission_answer.content)
                    if perm_answer_str.startswith(f"{self.custom_prefix}"):
                        await ctx.author.send(
                            "Παρακαλώ μη γράφετε εντολές κατά τη διαδικασία επιβεβαίωσης.Προσπαθήστε ξανά")

                if (perm_answer_str == "nai") or (perm_answer_str == "ναι") or (perm_answer_str == "yes"):
                    return True
                else:
                    return False

            # is the agent chat function turned on ? Είναι το agent chat ενεργοποιημένο?
            if self.bot.agent == False:
                return
            # is the user the bot or a real person ? //// Είναι ο χρήστης μποτ ή πραγματικό άτομο?
            if (ctx.author != self.bot.user) :
                # is the user participating in agent chat ?//// Ο χρήστης συμμετέχει ήδη σε συνομιλία με πράκτορα?
                if ctx.author.id in self.bot.agent_teams["ids"]:
                    # is this the channel the user is supposed to chat with the agent ? //// Είναι αυτό το κανάλι στο οποίο έχει ορισθεί η συνομιλία με το πράκτορα?
                    if ctx.channel.id == self.bot.agent_teams["ids"][ctx.author.id][0]:
                        # in which agent does the user participate ? //// Σε ποιο πράκτορα συμμετέχει ο χρήστης ?
                        name_of_agent = self.bot.agent_teams["ids"][ctx.author.id][1]
                        # check if a keyword the user mentioned is in the list of keywords of the agent the user participates.////Ψάξε λέξη κλειδί από τη συνομιλία συμμετεχόντων στο συγκεκριμένο κανάλι.
                        for keyword in self.bot.agent_questions[name_of_agent]:
                            # if you find a keyword , do the following : //// Αν βρεις λέξη κλειδί κάνε αυτά που έπονται :
                            if (ctx.content.find(str(keyword)) != -1):
                                keyword_questions = list(self.bot.agent_questions[name_of_agent][str(keyword)])
                                number_of_question = random.choice(keyword_questions)
                                #check what type of question we have sent so that we store it in a proper manner
                                type_of_question = int(self.bot.agent_questions[name_of_agent][str(keyword)][int(number_of_question)][0])

                                """ 
                                    Σε αυτό το σημείο με ένα if statement καταλαβαίνουμε τον τύπο της ερώτησης.Ο τύπος 0
                                    αναφέρεται σε ερωτήσεις αναπτύξεως. Ανάλογα το τύπο ερώτησης διαφορετικές ενέργειες 
                                    πρέπει να παρθούν για την αποθήκευση απαντήσεων.
                                """



                                if type_of_question == 0:



                                    """
                                        Σε αυτό το σημείο στέλνουμε το κυρίως σώμα της ερώτησης.
                                        Εδώ μπορούμε επίσης να διαλέξουμε σε ποιον θα απευθυνουμε την ερώτηση.
                                        ctx.channel.send σημαίνει ότι η ερώτηση απευθύνεται σε όλη την ομάδα.
                                        Μπορούμε να διαλέξουμε ένα από τα μέλη της ομάδας για παράδειγμα.
                                        Αλλά εδώ για λόγους ομαδικότητας και συνεργασίας απευθυνόμαστε σε
                                        όλα τα μέλη της ομάδας και αναθέτουμε την ευθύνη της απάντησης σε
                                        αυτόν που ενεργοποίησε τη λέξη κλειδί. Παράλληλα τα άλλα μέλη μπορούν να συζητούν
                                        μεταξύ τους στο ίδιο κανάλι . Το μποτ περιμένει απάντηση από αυτόν που ενεργοποίησε
                                        τη λέξη κλειδί.
                                        
                                        At this point we send the main body of the question
                                        Here we can also chose to whom the questions are asked.
                                        ctx.channel.send means the questions are asked to the team. 
                                        We can chose one of the team members for example.But here
                                        for the purpose of teamwork we ask the whole team.
                                    """


                                    await ctx.channel.send(
                                        f"```{str(self.bot.agent_questions[name_of_agent][str(keyword)][int(number_of_question)][1])}```\n"  # sending the main body of the question.
                                        f"Παρακαλώ γράψτε την απάντησή σας. ``Πατήστε Shift+enter για αλλαγή γραμμής`` . Υπεύθυνος απάντησης : **{ctx.author.name}**")
                                    self.bot.participating_in_agent_right_now.append(ctx.channel.id)
                                    this_is_the_final_answer = False
                                    while this_is_the_final_answer!=True:

                                        team_answer = await self.bot.wait_for("message", check=check)
                                        team_answer_str = str(team_answer.content)
                                        await ctx.channel.send(f"**Είναι αυτή η τελική σας απάντηση;**\n"
                                                               f"``{team_answer_str}``")
                                        this_is_the_final_answer = await permission(ctx)
                                        if this_is_the_final_answer!=True:
                                            await ctx.channel.send("Παρακαλώ γράψτε την απάντησή σας")
                                    self.bot.participating_in_agent_right_now.remove(ctx.channel.id)

                                    #here we need to store the answer
                                    #First create a profile for every user in the team if there isn't one.
                                    active_channel = self.bot.agent_teams["ids"][ctx.author.id][0]


                                    """
                                        Εδώ είναι ο τρόπος με τον οποίο οι απαντήσεις αποθηκεύονται.Ο παρακάτω κώδικας
                                        ερμηνεύεται ως εξής : Για όλα τα μέλη της ομάδας αυτού που απάντησε αποθήκευσε 
                                        στη μνήμη την απάντησή του ως απαντήσεις τους.
                                    """

                                    current_team = self.bot.agent_teams["teams"][active_channel]
                                    for id_element in current_team:
                                        if id_element not in self.bot.agent_answers:
                                            self.bot.agent_answers[id_element] = {}
                                        if name_of_agent not in self.bot.agent_answers[id_element]:
                                            self.bot.agent_answers[id_element][name_of_agent] = {}
                                        if str(keyword) not in self.bot.agent_answers[id_element][name_of_agent]:
                                            self.bot.agent_answers[id_element][name_of_agent][str(keyword)] = {}
                                        self.bot.agent_answers[id_element][name_of_agent][str(keyword)][number_of_question] = [type_of_question,team_answer_str]

                    else:
                        return
                else:
                    return
            else:
                return
        except: #we don't want to print the errors  created by every message !!! . We simply return !
            return
    """
        Η παρακάτω εντολή 'teamanswers' βασίζεται στο γεγονός ότι τα μέλη συμμετέχουν ακόμα στη συνομιλία με το 
        πράκτορα και οι ομάδες τους δεν έχουν χωρίσει ακόμα.Επίσης βασίζεται στην παραδοχή ότι οι απαντήσεις του πρώτου
        μέλους της ομάδας είναι επίσης οι απαντήσεις των υπόλοιπων μελών.Έτσι και αλλιώς τα καινούρια μέλη μιας ομάδας 
        προστίθονται στο τέλος της λίστας με append.Οπότε τα αρχικά μέλη είναι αυτά που έχουν απαντήσεις που 
        αντιπροσωπεύουν την ομάδα.
        
        !!!! this command 'teamanswers' is based on the hypothesis that the answers of the first person in a team are 
        also the answers of his teammates.It also requires members to still be participating in agent chat and their 
        teams are still formed and not broken apart.
    """ #!!! Teamanswers manual.
    @commands.command(name="teamanswers")
    async def teamanswers(self,ctx,name_of_agent=None):
        """
            Prints the answers of a team.
            <command_prefix>teamanswers <name_of_agent.docx>
        """
        try:
            if name_of_agent != None:
                if ctx.author.guild_permissions.administrator:
                    mycwd = os.getcwd()

                    name_of_agent_directory = name_of_agent.split('.')[0]

                    if not os.path.exists(f"{name_of_agent_directory} team answers"):
                        os.mkdir(f"{name_of_agent_directory} team answers")
                        await ctx.author.send(f"Folder '{name_of_agent_directory} team answers' δημιουργήθηκε στη τοποθεσία : {mycwd}")
                    else:
                        await ctx.author.send("Ο φάκελος υπάρχει ήδη , οι απαντήσεις αποθηκεύτηκαν.")
                    os.chdir(f"{name_of_agent_directory} team answers")
                    for channel_key in self.bot.agent_teams["teams"]:
                        #student_name here refers to student_id(gets the first student's id of the first team)
                        student_name = self.bot.agent_teams["teams"][channel_key][0]
                        if self.bot.agent_teams["ids"][student_name][1]==name_of_agent :
                            this_team_doc = docx.Document()
                            this_team_channel = self.bot.get_channel(channel_key)
                            this_team_doc.add_heading(f"Όνομα καναλιού ομάδας : {this_team_channel.name}\n"
                                                      f"Όνομα πράκτορα : {name_of_agent_directory}\n")

                            for student_id in self.bot.agent_teams["teams"][channel_key]:
                                if student_id in self.bot.agent_students:
                                    user_aem = self.bot.agent_students[student_id]["aem"]
                                    user_name = self.bot.agent_students[student_id]["name"]
                                    user_surname = self.bot.agent_students[student_id]["surname"]
                                    user_discord_id = student_id
                                    user_mail = self.bot.agent_students[student_id]["mail"]
                                    this_team_doc.add_paragraph(f"{user_aem} : {user_name} {user_surname} {user_mail} {user_discord_id} ")

                                else:
                                    this_team_doc.add_paragraph(f"Χωρίς εγγραφή : {student_id}")
                            for keyword in self.bot.agent_answers[student_name][name_of_agent]:
                                for question in self.bot.agent_answers[student_name][name_of_agent][keyword]:
                                    team_answer_body = self.bot.agent_answers[student_name][name_of_agent][keyword][question][1]
                                    type_of_question = self.bot.agent_answers[student_name][name_of_agent][keyword][question][0]
                                    if type_of_question==0 :
                                        this_team_doc.add_paragraph(f"{keyword}-{question}:\n{team_answer_body}")
                                    this_team_doc.add_paragraph(f"=============================================")
                            this_team_doc.save(f"{this_team_channel.name} {name_of_agent_directory}.docx")
                    os.chdir(f"{mycwd}")

                else:
                    await ctx.author.send("Δεν έχετε δικαιώματα διαχειριστή για αυτή την εντολή.")
            else:
                await ctx.author.send(f" ``Τα ορίσματα δεν προσφέρθηκαν σωστά``\n"
                                      f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}teamanswers <name_of_agent.docx>")
        except Exception as error:
            await ctx.author.send(f"{error}")

    @commands.command(name="agentanswers")
    async def agentanswers(self,ctx,name_of_agent=None):
        """
            Creates documents of student agent answers.
            <command_prefix>agentanswers <name_of_agent.docx>
        """
        try:
            if name_of_agent != None:
                if ctx.author.guild_permissions.administrator:
                    mycwd = os.getcwd()

                    name_of_agent_directory = name_of_agent.split('.')[0]

                    if not os.path.exists(f"{name_of_agent_directory} student agent answers"):
                        os.mkdir(f"{name_of_agent_directory} student agent answers")
                        await ctx.author.send(
                            f"Folder '{name_of_agent_directory} student agent answers' δημιουργήθηκε στη τοποθεσία : {mycwd}")
                    else:
                        await ctx.author.send("Ο φάκελος υπάρχει ήδη , οι απαντήσεις αποθηκεύτηκαν.")
                    os.chdir(f"{name_of_agent_directory} student agent answers")
                    #print(self.bot.agent_answers)
                    for student_id in self.bot.agent_answers:
                        if name_of_agent in self.bot.agent_answers[student_id]:
                            this_student_doc = docx.Document()
                            student_in_database = False
                            try:

                                user_aem = self.bot.agent_students[student_id]["aem"]
                                user_name = self.bot.agent_students[student_id]["name"]
                                user_surname = self.bot.agent_students[student_id]["surname"]
                                user_discord_id = student_id
                                user_mail = self.bot.agent_students[student_id]["mail"]
                                this_student_doc.add_heading(
                                    f"{user_aem} : {user_name} {user_surname} {user_mail} {user_discord_id} ")
                                student_in_database = True

                            except:
                                this_student_doc.add_heading(f"Χωρίς εγγραφή : {student_id}")

                            for keyword in self.bot.agent_answers[student_id][name_of_agent]:
                                for question in self.bot.agent_answers[student_id][name_of_agent][keyword]:
                                    student_answer_body = self.bot.agent_answers[student_id][name_of_agent][keyword][question][1]
                                    type_of_question =  self.bot.agent_answers[student_id][name_of_agent][keyword][question][0]
                                    if type_of_question == 0:
                                        this_student_doc.add_paragraph(f"{keyword}-{question}:\n{student_answer_body}")
                                    this_student_doc.add_paragraph(f"=============================================")
                            if student_in_database:
                                this_student_doc.save(f"{user_aem} {user_name} {user_surname} {name_of_agent_directory}.docx")
                            else:
                                this_student_doc.save(f"{student_id} {name_of_agent_directory}.docx")






                    os.chdir(f"{mycwd}")
                else:
                    await ctx.author.send("Δεν έχετε δικαιώματα διαχειριστή για αυτή την εντολή.")
            else:
                await ctx.author.send(f" ``Τα ορίσματα δεν προσφέρθηκαν σωστά``\n"
                                      f"**Σωστή μορφή ορισμάτων** : {self.custom_prefix}agentanswers <name_of_agent.docx>")
        except Exception as error:
            await ctx.author.send(f"{error}")

    @commands.command(name="agentparticipants")
    async def participants(self,ctx):
        try:
            if ctx.author.guild_permissions.administrator:
                dict_items = (self.bot.agent_teams["teams"]).items()
                string_to_send=""
                for element in dict_items:
                    element_channel = self.bot.get_channel(element[0])
                    string_to_send+=f"\n\n**{element_channel.name} : **"
                    this_team_users = ""
                    for r in range(len(element[1])):
                        element_user = self.bot.get_user(element[1][r])
                        this_team_users+=f"[{element_user.name}] "
                    string_to_send+=f"\n{this_team_users}"
                await ctx.author.send(string_to_send)
            else:
                await ctx.author.send("Δεν έχετε δικαιώματα διαχειριστή για αυτή την εντολή.")
        except Exception as error:
            await ctx.author.send(f"{error}")

    @commands.Cog.listener()
    async def on_ready(self):
        print('AgentChatCommands cog ready')





# Now, we need to set up this cog somehow, and we do that by making a setup function:
def setup(bot: commands.Bot):
    bot.add_cog(AgentChatCommands(bot))